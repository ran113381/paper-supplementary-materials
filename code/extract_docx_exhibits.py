from __future__ import annotations

import argparse
import csv
import re
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph
from lxml import etree


# Exhibit numbering follows the 2026-07 major revision (14 tables, 14 figures).
TABLE_SPECS = [
    ("table_01_sample_construction_procedure.csv", "Table 1 Sample Construction Procedure"),
    ("table_02_variable_definition_and_calculation.csv", "Table 2 Variable Definition and Calculation"),
    ("table_03_genai_dictionary_keyword_classification.csv", "Table 3 Classification of GenAI Dictionary Keywords Used for Text Mining"),
    ("table_04_full_sample_llm_semantic_audit.csv", "Table 4 Full-Sample LLM-Assisted Semantic Audit"),
    ("table_05_gold_standard_validation.csv", "Table 5 Human-Annotated Gold-Standard Validation of the Dictionary-Based Measure"),
    ("table_06_sota_hybrid_benchmark.csv", "Table 6 SOTA and Hybrid Benchmark Validation of GenAI Disclosure Measures"),
    ("table_07_descriptive_statistics.csv", "Table 7 Descriptive Statistics of Key Variables"),
    ("table_08_correlation_matrix.csv", "Table 8 The Correlation Matrix"),
    ("table_09_baseline_regression_results.csv", "Table 9 Baseline Regression Results"),
    ("table_10_robustness_summary.csv", "Table 10 Summary of Robustness Tests"),
    ("table_11_dyadic_dependence_corrections.csv", "Table 11 Robustness of Baseline Estimates to Dyadic-Dependence Corrections"),
    ("table_12_inventory_turnover_channel.csv", "Table 12 Candidate Operational Channel Test: Inventory Turnover"),
    ("table_13_fy2025_extension.csv", "Table 13 Extension of Baseline Estimates to Fiscal Year 2025"),
    ("table_14_heterogeneity_analysis.csv", "Table 14 Heterogeneity Analysis"),
]

# Captions are matched by prefix, because some manuscript captions carry a
# trailing explanatory sentence (e.g. Figure 11's scaling note).
FIGURE_SPECS = [
    ("figure_01", "Figure 1. Time evolution of GenAI disclosure rates in Chinese listed companies"),
    ("figure_02", "Figure 2. Theoretical Framework and Research Propositions"),
    ("figure_03", "Figure 3. Full-Sample LLM-Assisted Semantic Audit"),
    ("figure_04", "Figure 4. Dictionary Labels versus Strict LLM Consensus"),
    ("figure_05", "Figure 5. Differences in ROA Distribution Among Focal Enterprises by GenAI Disclosure Status"),
    ("figure_06", "Figure 6. R&D-Intensity Moderation Pattern"),
    ("figure_07", "Figure 7. Power Pressure Moderation Pattern"),
    ("figure_08", "Figure 8. Placebo Test: Distribution of Random-Permutation Coefficients"),
    ("figure_09", "Figure 9. Staggered DiD Event-Study Estimates"),
    ("figure_10", "Figure 10. Covariate Balance before and after Propensity Score Matching"),
    ("figure_11", "Figure 11. Candidate-Channel Path of Inventory Turnover"),
    ("figure_12", "Figure 12. SHAP-based Feature Importance: Full Sample versus GenAI Adopters"),
    ("figure_13", "Figure 13. SHAP Swarm Plot for the GenAI Adopter Subsample"),
    ("figure_14", "Figure 14. Heterogeneity Analysis: Forest Plot of Subsample Coefficients"),
]

# A caption paragraph is "Figure <n>. ..."; a body cross-reference is
# "Figure <n> presents ...". Only the former is an exhibit caption.
CAPTION_RE = re.compile(r"^Figure \d+\.")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Extract frozen tables and figures from the authority manuscript docx.")
    parser.add_argument("--docx", required=True, help="Path to the authority manuscript .docx")
    parser.add_argument("--table-dir", default="output/tables", help="Directory for exported table CSV files")
    parser.add_argument("--figure-dir", default="output/figures/manuscript_render", help="Directory for exported figure image files")
    return parser.parse_args()


def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def write_csv(path: Path, rows: list[list[str]]) -> None:
    with path.open("w", newline="", encoding="utf-8-sig") as fh:
        writer = csv.writer(fh)
        writer.writerows(rows)


def export_tables(doc: Document, out_dir: Path) -> None:
    if len(doc.tables) != len(TABLE_SPECS):
        raise RuntimeError(f"Expected {len(TABLE_SPECS)} tables, found {len(doc.tables)}")

    index_rows = [["file_name", "authority_label", "row_count", "column_count"]]
    for table, (filename, label) in zip(doc.tables, TABLE_SPECS):
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip().replace("\n", " / ") for cell in row.cells])
        write_csv(out_dir / filename, rows)
        index_rows.append([filename, label, str(len(rows)), str(len(rows[0]))])

    write_csv(out_dir / "TABLE_INDEX.csv", index_rows)


def read_body_image_rel_ids(docx_path: Path) -> tuple[list[str], dict[str, str]]:
    ns = {
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    }
    rel_ns = {
        "rel": "http://schemas.openxmlformats.org/package/2006/relationships",
    }

    with ZipFile(docx_path) as zf:
        document_xml = etree.fromstring(zf.read("word/document.xml"))
        rels_xml = etree.fromstring(zf.read("word/_rels/document.xml.rels"))
        rel_ids = [
            blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
            for blip in document_xml.xpath(".//a:blip", namespaces=ns)
        ]
        rel_map = {
            rel.get("Id"): rel.get("Target")
            for rel in rels_xml.xpath(".//rel:Relationship", namespaces=rel_ns)
        }
    return rel_ids, rel_map


def figure_captions(doc: Document) -> list[str]:
    captions: list[str] = []
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            paragraph = Paragraph(child, doc)
            text = paragraph.text.strip()
            if CAPTION_RE.match(text):
                captions.append(text)
    return captions


def export_figures(doc: Document, docx_path: Path, out_dir: Path) -> None:
    captions = figure_captions(doc)
    expected_captions = [caption for _, caption in FIGURE_SPECS]
    if len(captions) != len(expected_captions) or not all(
        found.startswith(expected) for found, expected in zip(captions, expected_captions)
    ):
        raise RuntimeError("Figure caption order does not match the expected manuscript structure")

    rel_ids, rel_map = read_body_image_rel_ids(docx_path)
    if len(rel_ids) != len(FIGURE_SPECS):
        raise RuntimeError(f"Expected {len(FIGURE_SPECS)} embedded figures, found {len(rel_ids)}")

    index_rows = [["file_name", "authority_caption", "source_media_name"]]
    with ZipFile(docx_path) as zf:
        for rel_id, (stem, caption) in zip(rel_ids, FIGURE_SPECS):
            target = rel_map[rel_id]
            source_name = Path(target).name
            ext = Path(source_name).suffix.lower()
            out_name = f"{stem}{ext}"
            with zf.open(f"word/{target}") as src, (out_dir / out_name).open("wb") as dst:
                dst.write(src.read())
            index_rows.append([out_name, caption, source_name])

    write_csv(out_dir / "FIGURE_INDEX.csv", index_rows)


def main() -> None:
    args = parse_args()
    docx_path = Path(args.docx).expanduser().resolve()
    table_dir = Path(args.table_dir)
    figure_dir = Path(args.figure_dir)

    if not docx_path.exists():
        raise FileNotFoundError(f"Manuscript not found: {docx_path}")

    ensure_dir(table_dir)
    ensure_dir(figure_dir)

    doc = Document(str(docx_path))
    export_tables(doc, table_dir)
    export_figures(doc, docx_path, figure_dir)

    print(f"Exported tables to: {table_dir}")
    print(f"Exported figures to: {figure_dir}")


if __name__ == "__main__":
    main()
